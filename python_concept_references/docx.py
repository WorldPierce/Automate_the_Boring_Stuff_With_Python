import docx # pip install python-docx

d = docx.Document('c:\\users\\wffpi\\documents\\demo.docx')

print(d.paragraphs)
print(d.paragrpahs[0].text)

p = d.paragraph[1]
print(p.runs[0].text) # runs are when text format changes
print(p.runs[1].bold) # returns true/false

p.runs[3].underline = True
p.runs[3].text = 'Underlined text'

d.save('c:\\demo2.docx')

p.style = 'Title' # can change styles

d = docx.Document() # new document not saved

d.add_paragraph('Hello, this is a paragraph')
d.add_paragraph('Paragraph number 2')
d.save('c:\\demo3.docx')

p = d.paragrphs[0]
p.add_run('New run!')
p.runs[1].bold = True
d.save('c:\\demo4.docx')


